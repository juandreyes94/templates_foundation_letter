"""
Document Generation Service
Generates Word documents from a template by replacing placeholders.
Deploy on Google Cloud Run - up to 2M free calls/month.
"""

import os
import io
import shutil
import tempfile
from flask import Flask, request, jsonify, send_file
from datetime import datetime
from docx import Document

app = Flask(__name__)

# Path to the template file (bundled with the container)
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

# All supported placeholders in the template
PLACEHOLDERS = [
    "date_format_july22-2025",
    "client_name",
    "Address_1",
    "Address_2",
    "Zip_code",
    "subdivision",
    "Project_Address",
    "Block",
    "Lot",
    "City",
    "print_date",
    "print_date_2",
    "IRC",
    "Soils_report_source",
    "Soils_report_number",
    "Soils_report_date_formatted_july9-2024",
]


def replace_in_paragraph(paragraph, replacements):
    """
    Replace placeholders in a paragraph handling Word's run fragmentation.
    Rebuilds the full paragraph text, replaces placeholders, then puts
    the result back in the first run (preserving its formatting).
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(("{" + key + "}") in full_text for key in replacements):
        return  # Nothing to replace in this paragraph

    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace("{" + key + "}", str(value) if value is not None else "")

    # Put the replaced text in the first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def generate_document(data: dict) -> bytes:
    """
    Takes a dict of field values, fills the template using python-docx,
    and returns the generated .docx file as bytes.
    """
    doc = Document(TEMPLATE_PATH)

    # Replace in all top-level paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)

    # Replace in headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, data)

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer.read()


@app.route("/health", methods=["GET"])
def health():
    """Health check endpoint."""
    return jsonify({"status": "ok", "template": os.path.exists(TEMPLATE_PATH)})


@app.route("/generate", methods=["POST"])
def generate():
    """
    Main endpoint. Accepts JSON body with field values.
    Returns the filled .docx file as a download.
    """
    if not request.is_json:
        return jsonify({"error": "Content-Type must be application/json"}), 400

    data = request.get_json()

    # Build replacement map — use empty string for missing fields
    replacements = {key: data.get(key, "") for key in PLACEHOLDERS}

    try:
        doc_bytes = generate_document(replacements)
    except Exception as e:
        return jsonify({"error": f"Document generation failed: {str(e)}"}), 500

    # Determine output filename
    filename = data.get("filename") or f"document_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"

    return send_file(
        io.BytesIO(doc_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/fields", methods=["GET"])
def fields():
    """Returns the list of all supported placeholder fields."""
    return jsonify({"fields": PLACEHOLDERS})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port, debug=False)
